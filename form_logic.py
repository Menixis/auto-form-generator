# -*- coding: utf-8 -*-
"""
Form Logic — pure logic for document generation (no UI dependencies)
=====================================================================
ใช้ร่วมกันได้ระหว่าง:
  - auto_form_generator.py  (Desktop GUI ผ่าน tkinter)
  - streamlit_app.py        (Web app ผ่าน Streamlit)
ไม่ import tkinter หรือ streamlit เลย ใช้แค่ python-docx
"""

import os
from docx import Document


# =============================================================================
# ค่าคงที่: ข้อมูลแบบฟอร์ม
# =============================================================================
FORM_TYPES = [
    {"id": 1, "filename": "form1.docx",
     "label": "1. หนังสือแจ้งให้เริ่มเข้าทำงาน / ส่งมอบผลิตภัณฑ์ (ครั้งแรก)",
     "uses_po": True, "uses_doc_type": True},
    {"id": 2, "filename": "form2.docx",
     "label": "2. หนังสือแจ้งให้กลับเข้าทำงาน (หลังแจ้งหยุดงาน)",
     "uses_po": False, "uses_doc_type": False},
    {"id": 3, "filename": "form3.docx",
     "label": "3. หนังสือแจ้งหยุดงาน (มิใช่ความผิดของ ปตท. และคู่สัญญา)",
     "uses_po": False, "uses_doc_type": False},
    {"id": 4, "filename": "form4.docx",
     "label": "4. หนังสือแจ้งหยุดงาน (เหตุจากความผิดของคู่สัญญา)",
     "uses_po": False, "uses_doc_type": False},
    {"id": 5, "filename": "form5.docx",
     "label": "5. หนังสือแจ้งหยุดงาน (เหตุจากความผิดของ ปตท.)",
     "uses_po": False, "uses_doc_type": False},
]

DOC_TYPE_OPTIONS = ["แจ้งให้เริ่มเข้าทำงาน", "แจ้งให้เริ่มส่งมอบสินค้า"]

# Placeholder ในแบบฟอร์มต้นฉบับ
# Shared across forms 1-5 (templates 2-5 were edited to use the same strings)
PLACEHOLDER_VENDOR = "Vendor"
PLACEHOLDER_WORK = "ชื่องาน"
PLACEHOLDER_PO = "PO_No"

# Form 1 only
PLACEHOLDER_DOC_NUM_F1 = "หมายเลขเอกสาร"
PLACEHOLDER_TYPE_F1 = "แจ้งให้เริ่มเข้าทำงาน / เริ่มส่งมอบสินค้า"
PLACEHOLDER_SIGNER_F1 = "ผู้มีอำนาจอนุมัติ หรือ ประธานกรรมการตรวจรับ"

# Forms 2-5 only (layout-based placeholders)
PLACEHOLDER_DOC_NUM_F2_5 = "                   /"
PLACEHOLDER_SIGNER_F2_5 = (
    "                                                                   "
    "(                                            )"
)

# Mapping จาก column ใน Excel → label ที่แสดงในผลลัพธ์
DISPLAY_FIELDS = [
    ("poNo", "เลข PO"),
    ("VendorName", "ชื่อคู่สัญญา"),
    ("title", "ชื่องาน"),
    ("prDocument", "เลขที่ PR"),
    ("procurementType", "ประเภทการจัดหา"),
    ("procurementMethod", "วิธีจัดหา"),
    ("prCreatorUnitName", "หน่วยงานผู้สร้าง PR"),
    ("budgetTH", "งบประมาณ (บาท)"),
    ("poCreateDate", "วันที่สร้าง PO"),
    ("cmsContractNo", "เลขที่สัญญา CMS"),
    ("daySLA", "จำนวนวัน SLA"),
]


# =============================================================================
# Utility: การแทนที่ข้อความใน docx โดยรักษารูปแบบเดิม
# =============================================================================
def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_text(doc, old_text, new_text):
    """แทนที่ข้อความทุกจุดใน document คืนค่าจำนวนครั้งที่แทนที่"""
    if not old_text:
        return 0
    count = 0
    for paragraph in _iter_paragraphs(doc):
        replaced = False
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                count += 1
                replaced = True
                break
        if replaced:
            continue
        if old_text in paragraph.text and paragraph.runs:
            new_full = paragraph.text.replace(old_text, new_text, 1)
            paragraph.runs[0].text = new_full
            for r in paragraph.runs[1:]:
                r.text = ""
            count += 1
    return count


def center_text_in_parens(name, total_width=44):
    inner = f" {name} ".strip()
    if len(inner) >= total_width:
        return f"({inner})"
    pad = total_width - len(inner)
    return "(" + " " * (pad // 2) + inner + " " * (pad - pad // 2) + ")"


def format_signer_line_f2_5(name):
    return " " * 67 + center_text_in_parens(name, total_width=44)


# =============================================================================
# Logic หลัก: สร้างเอกสารตามชนิดแบบฟอร์ม
# =============================================================================
def generate_form(form_id, template_path, output_path,
                  doc_number, po_number, doc_type, signer,
                  lookup_data=None):
    """
    สร้างเอกสาร Word จาก template
    Args:
        form_id: 1-5 (ประเภทแบบฟอร์ม)
        template_path: path ของไฟล์ template .docx
        output_path: path string หรือ BytesIO object สำหรับบันทึก
        doc_number: หมายเลขเอกสาร (เช่น PTT-001/2568)
        po_number: เลข PO
        doc_type: ประเภทเอกสาร (เฉพาะ form 1)
        signer: ชื่อผู้ลงนาม
        lookup_data: dict ของข้อมูลจาก Excel (สำหรับ auto-fill Vendor, ชื่องาน)
    Returns:
        list ของข้อความ log บอกว่าแทนที่อะไรไปบ้าง
    """
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"ไม่พบไฟล์ template: {template_path}")

    doc = Document(template_path)
    log = []
    vendor_name = (lookup_data or {}).get("VendorName") or ""
    work_title = (lookup_data or {}).get("title") or ""

    if form_id == 1:
        n = replace_text(doc, PLACEHOLDER_DOC_NUM_F1, doc_number or "")
        log.append(f"หมายเลขเอกสาร: แทนที่ {n} จุด")

        if doc_type:
            n = replace_text(doc, PLACEHOLDER_TYPE_F1, doc_type)
            log.append(f"ประเภทเอกสาร: แทนที่ {n} จุด")

        n = replace_text(doc, PLACEHOLDER_SIGNER_F1, signer or "")
        log.append(f"ผู้ลงนาม: แทนที่ {n} จุด")
    else:
        n = replace_text(
            doc, PLACEHOLDER_DOC_NUM_F2_5,
            f"  {doc_number}  " if doc_number else PLACEHOLDER_DOC_NUM_F2_5)
        log.append(f"หมายเลขเอกสาร: แทนที่ {n} จุด")

        if signer:
            n = replace_text(doc, PLACEHOLDER_SIGNER_F2_5,
                             format_signer_line_f2_5(signer))
            log.append(f"ผู้ลงนาม: แทนที่ {n} จุด")

    # Shared across all forms (templates 2-5 now use the same placeholder strings)
    if po_number:
        n = replace_text(doc, PLACEHOLDER_PO, po_number)
        log.append(f"เลข PO: แทนที่ {n} จุด")
    if vendor_name:
        n = replace_text(doc, PLACEHOLDER_VENDOR, vendor_name)
        log.append(f"ชื่อคู่สัญญา (Excel): แทนที่ {n} จุด")
    if work_title:
        n = replace_text(doc, PLACEHOLDER_WORK, work_title)
        log.append(f"ชื่องาน (Excel): แทนที่ {n} จุด")

    doc.save(output_path)
    return log
